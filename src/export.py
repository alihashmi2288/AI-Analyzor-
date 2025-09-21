"""
Export module for generating PDF and DOCX reports
Professional report generation with branding and formatting
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
from datetime import datetime
from typing import Dict, List
import streamlit as st

class ReportExporter:
    """Handles export of analysis results to various formats"""
    
    def __init__(self):
        self.company_name = "AI Resume Analyzer Pro"
        self.company_tagline = "Professional Resume Analysis & Optimization"
    
    def generate_pdf_report(self, results: Dict, user_info: Dict = None) -> bytes:
        """Generate comprehensive PDF report"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Build story (content)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#1f77b4')
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.HexColor('#666666')
            )
            
            # Header
            story.append(Paragraph(self.company_name, title_style))
            story.append(Paragraph(self.company_tagline, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Report title and date
            story.append(Paragraph("Resume Analysis Report", subtitle_style))
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", subtitle_style))
            
            summary_data = [
                ['Metric', 'Score', 'Status'],
                ['Overall Match', f"{results.get('similarity_score', 0)}%", self._get_status(results.get('similarity_score', 0))],
                ['ATS Score', f"{results.get('ats_score', 0)}/100", self._get_status(results.get('ats_score', 0))],
                ['Skills Match Rate', f"{results.get('skill_match_rate', 0)}%", self._get_status(results.get('skill_match_rate', 0))],
                ['Improvement Potential', f"{results.get('improvement_score', 0)}%", 'Moderate' if results.get('improvement_score', 0) > 50 else 'High']
            ]
            
            summary_table = Table(summary_data, colWidths=[2*inch, 1*inch, 1.5*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(summary_table)
            story.append(Spacer(1, 20))
            
            # Skills Analysis
            story.append(Paragraph("Skills Analysis", subtitle_style))
            
            matched_skills = results.get('matched_skills', [])
            missing_skills = results.get('missing_skills', [])
            
            if matched_skills:
                story.append(Paragraph("✅ Matched Skills:", styles['Heading3']))
                for i, skill in enumerate(matched_skills[:15], 1):
                    story.append(Paragraph(f"{i}. {skill}", styles['Normal']))
                story.append(Spacer(1, 10))
            
            if missing_skills:
                story.append(Paragraph("❌ Missing Skills (Recommendations):", styles['Heading3']))
                for i, skill in enumerate(missing_skills[:15], 1):
                    story.append(Paragraph(f"{i}. {skill}", styles['Normal']))
                story.append(Spacer(1, 10))
            
            # Keywords Analysis
            if results.get('resume_keywords'):
                story.append(Paragraph("Top Resume Keywords", subtitle_style))
                keywords_text = ", ".join([kw[0] for kw in results['resume_keywords'][:20]])
                story.append(Paragraph(keywords_text, styles['Normal']))
                story.append(Spacer(1, 20))
            
            # AI Suggestions
            if results.get('ai_suggestions'):
                story.append(Paragraph("AI-Powered Recommendations", subtitle_style))
                for i, suggestion in enumerate(results['ai_suggestions'][:10], 1):
                    story.append(Paragraph(f"{i}. {suggestion}", styles['Normal']))
                story.append(Spacer(1, 10))
            
            # ATS Feedback
            if results.get('ats_feedback'):
                story.append(Paragraph("ATS Optimization Tips", subtitle_style))
                for i, feedback in enumerate(results['ats_feedback'], 1):
                    story.append(Paragraph(f"{i}. {feedback}", styles['Normal']))
                story.append(Spacer(1, 10))
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph("This report was generated by AI Resume Analyzer Pro", styles['Normal']))
            story.append(Paragraph("For more information, visit our website or contact support.", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            st.error(f"Error generating PDF report: {str(e)}")
            return b""
    
    def generate_docx_report(self, results: Dict, user_info: Dict = None) -> bytes:
        """Generate DOCX report"""
        try:
            doc = Document()
            
            # Header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"{self.company_name} - {self.company_tagline}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Title
            title = doc.add_heading('Resume Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            
            summary_table = doc.add_table(rows=5, cols=3)
            summary_table.style = 'Table Grid'
            
            # Table headers
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Score'
            hdr_cells[2].text = 'Status'
            
            # Table data
            metrics = [
                ('Overall Match', f"{results.get('similarity_score', 0)}%", self._get_status(results.get('similarity_score', 0))),
                ('ATS Score', f"{results.get('ats_score', 0)}/100", self._get_status(results.get('ats_score', 0))),
                ('Skills Match Rate', f"{results.get('skill_match_rate', 0)}%", self._get_status(results.get('skill_match_rate', 0))),
                ('Improvement Potential', f"{results.get('improvement_score', 0)}%", 'Moderate' if results.get('improvement_score', 0) > 50 else 'High')
            ]
            
            for i, (metric, score, status) in enumerate(metrics, 1):
                row_cells = summary_table.rows[i].cells
                row_cells[0].text = metric
                row_cells[1].text = score
                row_cells[2].text = status
            
            # Skills Analysis
            doc.add_heading('Skills Analysis', level=1)
            
            matched_skills = results.get('matched_skills', [])
            missing_skills = results.get('missing_skills', [])
            
            if matched_skills:
                doc.add_heading('✅ Matched Skills', level=2)
                for skill in matched_skills[:15]:
                    doc.add_paragraph(f"• {skill}", style='List Bullet')
            
            if missing_skills:
                doc.add_heading('❌ Missing Skills (Recommendations)', level=2)
                for skill in missing_skills[:15]:
                    doc.add_paragraph(f"• {skill}", style='List Bullet')
            
            # Keywords Analysis
            if results.get('resume_keywords'):
                doc.add_heading('Top Resume Keywords', level=1)
                keywords_text = ", ".join([kw[0] for kw in results['resume_keywords'][:20]])
                doc.add_paragraph(keywords_text)
            
            # AI Suggestions
            if results.get('ai_suggestions'):
                doc.add_heading('AI-Powered Recommendations', level=1)
                for suggestion in results['ai_suggestions'][:10]:
                    doc.add_paragraph(f"• {suggestion}", style='List Bullet')
            
            # ATS Feedback
            if results.get('ats_feedback'):
                doc.add_heading('ATS Optimization Tips', level=1)
                for feedback in results['ats_feedback']:
                    doc.add_paragraph(f"• {feedback}", style='List Bullet')
            
            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph("This report was generated by AI Resume Analyzer Pro")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            st.error(f"Error generating DOCX report: {str(e)}")
            return b""
    
    def generate_cover_letter_docx(self, cover_letter_text: str, job_info: Dict = None) -> bytes:
        """Generate formatted cover letter in DOCX"""
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Date
            date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add space
            doc.add_paragraph()
            
            # Job info (if provided)
            if job_info:
                if job_info.get('company'):
                    doc.add_paragraph(f"Hiring Manager")
                    doc.add_paragraph(f"{job_info['company']}")
                    doc.add_paragraph()
            
            # Salutation
            doc.add_paragraph("Dear Hiring Manager,")
            doc.add_paragraph()
            
            # Cover letter content
            paragraphs = cover_letter_text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    para = doc.add_paragraph(paragraph.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Closing
            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph("[Your Name]")
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            st.error(f"Error generating cover letter: {str(e)}")
            return b""
    
    def _get_status(self, score: float) -> str:
        """Get status based on score"""
        if score >= 80:
            return "Excellent"
        elif score >= 60:
            return "Good"
        elif score >= 40:
            return "Fair"
        else:
            return "Needs Improvement"
    
    def generate_comparison_report(self, comparison_results: List[Dict]) -> bytes:
        """Generate multi-resume comparison report"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#1f77b4')
            )
            
            story.append(Paragraph("Multi-Resume Comparison Report", title_style))
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Comparison table
            table_data = [['Resume', 'Match Score', 'ATS Score', 'Skills Matched', 'Skills Missing']]
            
            for result in comparison_results:
                table_data.append([
                    result.get('filename', 'Unknown'),
                    f"{result.get('similarity_score', 0)}%",
                    f"{result.get('ats_score', 0)}/100",
                    str(len(result.get('matched_skills', []))),
                    str(len(result.get('missing_skills', [])))
                ])
            
            comparison_table = Table(table_data, colWidths=[2*inch, 1*inch, 1*inch, 1*inch, 1*inch])
            comparison_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(comparison_table)
            story.append(Spacer(1, 20))
            
            # Recommendations
            story.append(Paragraph("Recommendations", styles['Heading2']))
            
            # Find best performing resume
            best_resume = max(comparison_results, key=lambda x: x.get('similarity_score', 0))
            story.append(Paragraph(f"Best performing resume: {best_resume.get('filename', 'Unknown')} with {best_resume.get('similarity_score', 0)}% match", styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            st.error(f"Error generating comparison report: {str(e)}")
            return b""