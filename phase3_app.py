"""
Phase 3 - Pro-Level Features: SaaS-Style Resume Analyzer
Complete professional tool with authentication, branding, and advanced features
"""

import streamlit as st
import streamlit_authenticator as stauth
import sqlite3
import hashlib
import json
from datetime import datetime
import pdfplumber
import docx2txt
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import openai
import plotly.graph_objects as go
import plotly.express as px
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import io
import os
import base64
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Page config with branding
st.set_page_config(
    page_title="ResumeAI Pro",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional branding
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .feature-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
    }
    .stButton > button {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 5px;
        padding: 0.5rem 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Database setup
def init_database():
    conn = sqlite3.connect('resumeai.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY,
            username TEXT UNIQUE,
            email TEXT UNIQUE,
            password TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sessions (
            id INTEGER PRIMARY KEY,
            user_id INTEGER,
            resume_text TEXT,
            jd_text TEXT,
            match_score REAL,
            ats_score REAL,
            cover_letter TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database
init_database()

# Authentication functions
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def create_user(username, email, password):
    conn = sqlite3.connect('resumeai.db')
    cursor = conn.cursor()
    
    try:
        hashed_pw = hash_password(password)
        cursor.execute('INSERT INTO users (username, email, password) VALUES (?, ?, ?)',
                      (username, email, hashed_pw))
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        conn.close()
        return False

def authenticate_user(username, password):
    conn = sqlite3.connect('resumeai.db')
    cursor = conn.cursor()
    
    hashed_pw = hash_password(password)
    cursor.execute('SELECT id, username, email FROM users WHERE username = ? AND password = ?',
                  (username, hashed_pw))
    user = cursor.fetchone()
    conn.close()
    
    return user

def save_session(user_id, resume_text, jd_text, match_score, ats_score, cover_letter=""):
    conn = sqlite3.connect('resumeai.db')
    cursor = conn.cursor()
    
    cursor.execute('''INSERT INTO sessions 
                     (user_id, resume_text, jd_text, match_score, ats_score, cover_letter)
                     VALUES (?, ?, ?, ?, ?, ?)''',
                  (user_id, resume_text, jd_text, match_score, ats_score, cover_letter))
    conn.commit()
    conn.close()

def get_user_sessions(user_id):
    conn = sqlite3.connect('resumeai.db')
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM sessions WHERE user_id = ? ORDER BY created_at DESC',
                  (user_id,))
    sessions = cursor.fetchall()
    conn.close()
    
    return sessions

# Core functions from previous phases
def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(uploaded_file):
    return docx2txt.process(uploaded_file)

def clean_text(text):
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s.,;:!?()-]', ' ', text)
    return text.strip()

def calculate_match_score(resume_text, jd_text):
    if not resume_text or not jd_text:
        return 0
    
    clean_resume = clean_text(resume_text)
    clean_jd = clean_text(jd_text)
    
    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
    
    try:
        tfidf_matrix = vectorizer.fit_transform([clean_resume, clean_jd])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        return round(similarity[0][0] * 100, 1)
    except:
        return 0

def calculate_ats_score(resume_text, jd_text):
    if not resume_text or not jd_text:
        return 0
    
    jd_words = set(clean_text(jd_text).split())
    resume_words = set(clean_text(resume_text).split())
    
    overlap = len(jd_words & resume_words)
    total_jd_words = len(jd_words)
    
    if total_jd_words == 0:
        return 0
    
    ats_score = (overlap / total_jd_words) * 100
    
    action_verbs = ['managed', 'led', 'developed', 'created', 'improved', 'achieved']
    action_count = sum(1 for verb in action_verbs if verb in resume_text.lower())
    ats_score += min(action_count * 5, 20)
    
    return min(round(ats_score, 1), 100)

# Cover letter templates
COVER_LETTER_TEMPLATES = {
    "Formal": {
        "tone": "formal and professional",
        "structure": "traditional business letter format",
        "language": "conservative and respectful"
    },
    "Modern": {
        "tone": "contemporary and engaging",
        "structure": "modern format with personality",
        "language": "confident and dynamic"
    },
    "Creative": {
        "tone": "creative and innovative",
        "structure": "unique format showing creativity",
        "language": "expressive and original"
    },
    "Short": {
        "tone": "concise and direct",
        "structure": "brief 2-paragraph format",
        "language": "clear and to-the-point"
    }
}

def generate_cover_letter(resume_text, jd_text, job_title, template_type, api_key):
    if not api_key:
        return "Please provide OpenAI API key."
    
    try:
        openai.api_key = api_key
        template = COVER_LETTER_TEMPLATES[template_type]
        
        prompt = f"""
        Write a {template['tone']} cover letter using {template['structure']} with {template['language']} language.
        
        Job Title: {job_title}
        Resume: {resume_text[:1500]}
        Job Description: {jd_text[:1000]}
        
        Template Style: {template_type}
        - Tone: {template['tone']}
        - Structure: {template['structure']}
        - Language: {template['language']}
        
        Create a compelling cover letter that matches this style perfectly.
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        return f"Error: {str(e)}"

def generate_resume_rewrite(resume_text, jd_text, api_key):
    if not api_key:
        return "Please provide OpenAI API key."
    
    try:
        openai.api_key = api_key
        
        prompt = f"""
        Rewrite this resume to perfectly match the job description. Make it ATS-friendly and impactful.
        
        Original Resume: {resume_text[:2000]}
        Target Job: {jd_text[:1000]}
        
        Instructions:
        1. Keep the same structure but improve content
        2. Add relevant keywords from job description
        3. Quantify achievements where possible
        4. Use strong action verbs
        5. Tailor experience descriptions to match job requirements
        
        Return the complete rewritten resume.
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.6
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        return f"Error: {str(e)}"

def generate_interview_questions(resume_text, jd_text, api_key):
    if not api_key:
        return "Please provide OpenAI API key."
    
    try:
        openai.api_key = api_key
        
        prompt = f"""
        Generate 8 interview questions based on this resume and job description:
        
        Resume: {resume_text[:1500]}
        Job Description: {jd_text[:1000]}
        
        Create:
        - 4 technical/role-specific questions
        - 4 behavioral questions
        
        For each question, provide:
        1. The question
        2. Key points to address in answer
        3. Example answer framework
        
        Format as numbered list with clear sections.
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        return f"Error: {str(e)}"

def create_branded_pdf_report(analysis_data, cover_letter=""):
    pdf = FPDF()
    pdf.add_page()
    
    # Header with branding
    pdf.set_fill_color(102, 126, 234)
    pdf.rect(0, 0, 210, 30, 'F')
    
    pdf.set_font('Arial', 'B', 20)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 20, 'ResumeAI Pro - Analysis Report', 0, 1, 'C')
    
    pdf.set_text_color(0, 0, 0)
    pdf.ln(10)
    
    # Analysis results
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, 'Analysis Summary', 0, 1)
    pdf.ln(5)
    
    pdf.set_font('Arial', '', 12)
    pdf.cell(0, 8, f'Match Score: {analysis_data.get("match_score", 0)}%', 0, 1)
    pdf.cell(0, 8, f'ATS Score: {analysis_data.get("ats_score", 0)}%', 0, 1)
    pdf.cell(0, 8, f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', 0, 1)
    
    if cover_letter:
        pdf.ln(10)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Generated Cover Letter', 0, 1)
        pdf.ln(5)
        
        pdf.set_font('Arial', '', 10)
        lines = cover_letter.split('\n')
        for line in lines:
            if line.strip():
                pdf.cell(0, 5, line[:80], 0, 1)
    
    return pdf.output(dest='S').encode('latin-1')

def main():
    # Initialize session state
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'user_data' not in st.session_state:
        st.session_state.user_data = None
    
    # Header with branding
    st.markdown("""
    <div class="main-header">
        <h1>🎯 ResumeAI Pro</h1>
        <p>Professional AI-Powered Resume Analysis & Optimization</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Authentication
    if not st.session_state.authenticated:
        show_auth_page()
        return
    
    # Main application
    show_main_app()

def show_auth_page():
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("### Welcome to ResumeAI Pro")
        
        tab1, tab2 = st.tabs(["Sign In", "Sign Up"])
        
        with tab1:
            with st.form("login_form"):
                username = st.text_input("Username")
                password = st.text_input("Password", type="password")
                
                if st.form_submit_button("Sign In", use_container_width=True):
                    user = authenticate_user(username, password)
                    if user:
                        st.session_state.authenticated = True
                        st.session_state.user_data = {
                            'id': user[0],
                            'username': user[1],
                            'email': user[2]
                        }
                        st.rerun()
                    else:
                        st.error("Invalid credentials")
        
        with tab2:
            with st.form("signup_form"):
                new_username = st.text_input("Choose Username")
                new_email = st.text_input("Email")
                new_password = st.text_input("Password", type="password")
                confirm_password = st.text_input("Confirm Password", type="password")
                
                if st.form_submit_button("Sign Up", use_container_width=True):
                    if new_password != confirm_password:
                        st.error("Passwords don't match")
                    elif len(new_password) < 6:
                        st.error("Password must be at least 6 characters")
                    elif create_user(new_username, new_email, new_password):
                        st.success("Account created! Please sign in.")
                    else:
                        st.error("Username or email already exists")

def show_main_app():
    # Sidebar
    with st.sidebar:
        st.markdown(f"### Welcome, {st.session_state.user_data['username']}!")
        
        if st.button("🚪 Logout"):
            st.session_state.authenticated = False
            st.session_state.user_data = None
            st.rerun()
        
        st.divider()
        
        st.header("⚙️ Settings")
        api_key = st.text_input("OpenAI API Key", type="password")
        
        st.header("🎮 Quick Demo")
        if st.button("Load Sample Data"):
            # Load demo data
            st.session_state.demo_mode = True
            st.success("Demo data loaded!")
    
    # Main tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📊 Dashboard",
        "📄 Resume Rewrite", 
        "📝 Cover Letters",
        "❓ Interview Prep",
        "📈 Export & Reports",
        "📚 History"
    ])
    
    with tab1:
        show_dashboard_tab(api_key)
    
    with tab2:
        show_resume_rewrite_tab(api_key)
    
    with tab3:
        show_cover_letter_tab(api_key)
    
    with tab4:
        show_interview_prep_tab(api_key)
    
    with tab5:
        show_export_tab()
    
    with tab6:
        show_history_tab()

def show_dashboard_tab(api_key):
    st.header("📊 Resume Analysis Dashboard")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📄 Upload Resume")
        resume_file = st.file_uploader("Choose resume file", type=['pdf', 'docx'])
        
        if resume_file:
            if resume_file.type == "application/pdf":
                resume_text = extract_text_from_pdf(resume_file)
            else:
                resume_text = extract_text_from_docx(resume_file)
            
            st.session_state.resume_text = resume_text
            st.success(f"✅ {resume_file.name} uploaded")
    
    with col2:
        st.subheader("📋 Job Description")
        jd_input = st.text_area("Paste job description:", height=200)
        
        if jd_input:
            st.session_state.jd_text = jd_input
    
    # Analysis
    if st.button("🚀 Analyze Resume", type="primary", use_container_width=True):
        if hasattr(st.session_state, 'resume_text') and hasattr(st.session_state, 'jd_text'):
            with st.spinner("🔄 Analyzing your resume..."):
                match_score = calculate_match_score(st.session_state.resume_text, st.session_state.jd_text)
                ats_score = calculate_ats_score(st.session_state.resume_text, st.session_state.jd_text)
                
                st.session_state.match_score = match_score
                st.session_state.ats_score = ats_score
                
                # Save session
                if st.session_state.user_data:
                    save_session(
                        st.session_state.user_data['id'],
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        match_score,
                        ats_score
                    )
            
            # Display results
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown(f"""
                <div class="metric-card">
                    <h2>{match_score}%</h2>
                    <p>Match Score</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="metric-card">
                    <h2>{ats_score}%</h2>
                    <p>ATS Score</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                overall = (match_score + ats_score) / 2
                st.markdown(f"""
                <div class="metric-card">
                    <h2>{overall:.1f}%</h2>
                    <p>Overall</p>
                </div>
                """, unsafe_allow_html=True)
            
            # Gauge visualization
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=match_score,
                title={'text': "Resume Match Score"},
                gauge={'axis': {'range': [None, 100]},
                       'bar': {'color': "#667eea"},
                       'steps': [{'range': [0, 50], 'color': "lightgray"},
                               {'range': [50, 80], 'color': "yellow"},
                               {'range': [80, 100], 'color': "green"}]}
            ))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.error("Please upload resume and add job description")

def show_resume_rewrite_tab(api_key):
    st.header("📄 One-Click Resume Rewrite")
    
    if hasattr(st.session_state, 'resume_text') and hasattr(st.session_state, 'jd_text'):
        st.info("🎯 AI will rewrite your resume to perfectly match the job description")
        
        if st.button("✨ Rewrite My Resume", type="primary"):
            if api_key:
                with st.spinner("🤖 AI is rewriting your resume..."):
                    rewritten_resume = generate_resume_rewrite(
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        api_key
                    )
                
                st.subheader("📝 Your Optimized Resume")
                
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.markdown("**Original Resume:**")
                    st.text_area("Original", st.session_state.resume_text[:1000] + "...", height=300, disabled=True)
                
                with col2:
                    st.markdown("**AI-Optimized Resume:**")
                    edited_resume = st.text_area("Optimized", rewritten_resume, height=300)
                
                # Download options
                st.download_button(
                    "📥 Download Optimized Resume",
                    edited_resume,
                    "optimized_resume.txt",
                    "text/plain",
                    use_container_width=True
                )
            else:
                st.warning("🔑 Please add your OpenAI API key in the sidebar")
    else:
        st.info("📄 Please analyze a resume first in the Dashboard tab")

def show_cover_letter_tab(api_key):
    st.header("📝 AI Cover Letter Generator")
    
    if hasattr(st.session_state, 'resume_text') and hasattr(st.session_state, 'jd_text'):
        col1, col2 = st.columns([1, 1])
        
        with col1:
            job_title = st.text_input("Job Title", placeholder="e.g., Senior Software Engineer")
            template_type = st.selectbox("Cover Letter Style", list(COVER_LETTER_TEMPLATES.keys()))
        
        with col2:
            st.markdown("**Template Preview:**")
            template_info = COVER_LETTER_TEMPLATES[template_type]
            st.info(f"""
            **Tone:** {template_info['tone']}
            **Structure:** {template_info['structure']}
            **Language:** {template_info['language']}
            """)
        
        if st.button("✨ Generate Cover Letter", type="primary"):
            if api_key:
                with st.spinner(f"🤖 Creating {template_type.lower()} cover letter..."):
                    cover_letter = generate_cover_letter(
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        job_title,
                        template_type,
                        api_key
                    )
                
                st.subheader(f"📄 Your {template_type} Cover Letter")
                edited_letter = st.text_area("Edit your cover letter:", cover_letter, height=400)
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        "📝 Download as Text",
                        edited_letter,
                        f"cover_letter_{template_type.lower()}.txt",
                        "text/plain"
                    )
                
                with col2:
                    # Create DOCX
                    doc = Document()
                    doc.add_heading('Cover Letter', 0)
                    doc.add_paragraph(edited_letter)
                    
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    st.download_button(
                        "📄 Download as DOCX",
                        docx_buffer.getvalue(),
                        f"cover_letter_{template_type.lower()}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("🔑 Please add your OpenAI API key")
    else:
        st.info("📄 Please analyze a resume first")

def show_interview_prep_tab(api_key):
    st.header("❓ AI Interview Preparation")
    
    if hasattr(st.session_state, 'resume_text') and hasattr(st.session_state, 'jd_text'):
        st.info("🎯 Get personalized interview questions based on your resume and target job")
        
        if st.button("🎤 Generate Interview Questions", type="primary"):
            if api_key:
                with st.spinner("🤖 Preparing your interview questions..."):
                    questions = generate_interview_questions(
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        api_key
                    )
                
                st.subheader("📋 Your Personalized Interview Questions")
                st.markdown(questions)
                
                st.download_button(
                    "📥 Download Interview Prep",
                    questions,
                    "interview_questions.txt",
                    "text/plain",
                    use_container_width=True
                )
            else:
                st.warning("🔑 Please add your OpenAI API key")
    else:
        st.info("📄 Please analyze a resume first")

def show_export_tab():
    st.header("📈 Export & Reports")
    
    if hasattr(st.session_state, 'match_score'):
        st.subheader("📊 Available Reports")
        
        analysis_data = {
            'match_score': st.session_state.get('match_score', 0),
            'ats_score': st.session_state.get('ats_score', 0)
        }
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📄 Generate PDF Report"):
                pdf_data = create_branded_pdf_report(analysis_data)
                st.download_button(
                    "📥 Download PDF",
                    pdf_data,
                    "resume_analysis_report.pdf",
                    "application/pdf"
                )
        
        with col2:
            if st.button("📊 Export Data (JSON)"):
                json_data = json.dumps(analysis_data, indent=2)
                st.download_button(
                    "📥 Download JSON",
                    json_data,
                    "analysis_data.json",
                    "application/json"
                )
        
        with col3:
            if st.button("📧 Email-Ready Format"):
                email_text = f"""
Resume Analysis Summary:
- Match Score: {analysis_data['match_score']}%
- ATS Score: {analysis_data['ats_score']}%
- Generated: {datetime.now().strftime('%Y-%m-%d')}

Generated by ResumeAI Pro
                """
                st.download_button(
                    "📥 Download Email",
                    email_text,
                    "email_summary.txt",
                    "text/plain"
                )
    else:
        st.info("📄 Please analyze a resume first to generate reports")

def show_history_tab():
    st.header("📚 Analysis History")
    
    if st.session_state.user_data:
        sessions = get_user_sessions(st.session_state.user_data['id'])
        
        if sessions:
            st.subheader(f"📊 Your Past {len(sessions)} Analyses")
            
            for session in sessions:
                with st.expander(f"Analysis from {session[7][:16]} - Score: {session[4]}%"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.metric("Match Score", f"{session[4]}%")
                        st.metric("ATS Score", f"{session[5]}%")
                    
                    with col2:
                        if st.button(f"Load Session", key=f"load_{session[0]}"):
                            st.session_state.resume_text = session[2]
                            st.session_state.jd_text = session[3]
                            st.session_state.match_score = session[4]
                            st.session_state.ats_score = session[5]
                            st.success("Session loaded!")
        else:
            st.info("📝 No analysis history yet. Start by analyzing your first resume!")

if __name__ == "__main__":
    main()